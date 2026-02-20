"""
Pipeline de ingestão de documentos: PDF, Excel, URLs.
Extende execution/document_ingest.py para suportar mais formatos.
"""
import os
import uuid
from pathlib import Path
from typing import List, Optional
import requests
from bs4 import BeautifulSoup

CHUNK_SIZE = 600
CHUNK_OVERLAP = 80


def _extract_text_from_file(file_path: str) -> str:
    """Extrai texto de qualquer formato suportado."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
    
    suf = path.suffix.lower()
    
    if suf == ".txt":
        return path.read_text(encoding="utf-8", errors="replace").strip()
    
    if suf == ".pdf":
        return _extract_text_from_pdf(path)
    
    if suf in (".xlsx", ".xls"):
        return _extract_text_from_excel(path)
    
    if suf == ".docx":
        return _extract_text_from_docx(path)
    
    if suf == ".csv":
        return _extract_text_from_csv(path)
    
    if suf == ".md":
        return _extract_text_from_markdown(path)
    
    if suf == ".html":
        return _extract_text_from_html(path)
    
    raise ValueError(f"Formato não suportado: {suf}. Use .txt, .pdf, .xlsx, .xls, .docx, .csv, .md, .html.")


def _extract_text_from_pdf(path: Path) -> str:
    """Extrai texto de PDF usando pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("Para PDF instale: pip install pypdf")
    
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts).strip()


def _extract_text_from_excel(path: Path) -> str:
    """Extrai texto de Excel (.xlsx, .xls)."""
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("Para Excel instale: pip install openpyxl")
    
    parts = []
    wb = openpyxl.load_workbook(path, data_only=True)
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"\n--- Sheet: {sheet_name} ---\n")
        
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                parts.append(row_text)
    
    return "\n".join(parts).strip()


def _extract_text_from_docx(path: Path) -> str:
    """Extrai texto de Word (.docx)."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Para Word instale: pip install python-docx")
    
    doc = Document(str(path))
    parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    
    # Também extrai texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    
    return "\n".join(parts).strip()


def _extract_text_from_csv(path: Path) -> str:
    """Extrai texto de CSV."""
    import csv
    
    parts = []
    with open(path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = " | ".join(cell for cell in row)
            if row_text.strip():
                parts.append(row_text)
    
    return "\n".join(parts).strip()


def _extract_text_from_markdown(path: Path) -> str:
    """Extrai texto de Markdown (remove sintaxe, mantém conteúdo)."""
    import re
    
    text = path.read_text(encoding='utf-8', errors='replace')
    
    # Remove código
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'`[^`]+`', '', text)
    
    # Remove links mas mantém texto
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # Remove imagens
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    
    # Remove headers markers
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    
    # Remove bold/italic
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # Remove horizontais
    text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
    
    return text.strip()


def _extract_text_from_html(path: Path) -> str:
    """Extrai texto de HTML."""
    html = path.read_text(encoding='utf-8', errors='replace')
    return _extract_text_from_html_string(html)


def _extract_text_from_url(url: str, timeout: int = 30) -> str:
    """
    Extrai texto de uma URL (web scraping).
    Suporta HTML plain e renderizado (simples).
    """
    try:
        response = requests.get(url, timeout=timeout, headers={
            'User-Agent': 'Mozilla/5.0 (compatible; B&B-RAG-Bot/1.0)'
        })
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', '')
        
        if 'text/html' in content_type:
            return _extract_text_from_html_string(response.text)
        elif 'text/plain' in content_type:
            return response.text.strip()
        else:
            raise ValueError(f"Content-Type não suportado: {content_type}")
            
    except requests.RequestException as e:
        raise RuntimeError(f"Erro ao buscar URL: {e}")


def _extract_text_from_html_string(html: str) -> str:
    """Extrai texto de uma string HTML."""
    soup = BeautifulSoup(html, 'html.parser')
    
    # Remove scripts e styles
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    
    # Extrai texto
    text = soup.get_text(separator='\n')
    
    # Limpa whitespace
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]
    
    return '\n'.join(lines)


def _chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Divide texto em blocos com overlap."""
    if not text or not text.strip():
        return []
    
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + size
        chunk = text[start:end]
        
        if not chunk.strip():
            start = end - overlap
            continue
            
        chunks.append(chunk.strip())
        start = end - overlap
    
    return chunks


def get_file_size_mb(file_path: str) -> float:
    """Retorna tamanho do arquivo em MB."""
    return os.path.getsize(file_path) / (1024 * 1024)


def get_url_content_length(url: str) -> int:
    """Retorna tamanho do conteúdo da URL em bytes (HEAD request)."""
    try:
        response = requests.head(url, timeout=10, allow_redirects=True)
        return int(response.headers.get('content-length', 0))
    except:
        return 0
